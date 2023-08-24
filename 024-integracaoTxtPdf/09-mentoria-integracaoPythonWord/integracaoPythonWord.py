"""
Integração Python Word

Instalar biblioteca python-docx

pip install python-docx

"""

from docx import Document

# criando o documento no Python
documento = Document()

faturamento = 1000

# aqui você edita tudo o que você quer
texto = f"""Fala Lira,

O faturamento da empresa ontem foi de R${faturamento}

Tamo junto, abs.
"""

paragrafo = documento.add_paragraph(texto) 




### formatação

# Pt vem de pontos que seria o tamanho da fonte
# RGBColor é das cores
# Cm é de centimetros para editar tamanho de imagem, tamanho de margem, etc
from docx.shared import Pt, RGBColor, Cm # valores de formatação

# permite criar um estilo de formatação (determina um conjunto de características para a fonte como cor, tipo de fonte, tamanho, etc)
from docx.enum.style import WD_STYLE_TYPE

# Criamos um estilo para ser utilizado (conjunto de caracteristicas do texto)
# 'EstiloLira' é o nome do estilo que está sendo criado
paragrafo.style = documento.styles.add_style("EstiloLira", WD_STYLE_TYPE.PARAGRAPH)

paragrafo.style.font.name = "Algerian" # nome da fonte como está no word
paragrafo.style.font.size = Pt(15) # tamanho da fonte no padrao do word
paragrafo.style.font.bold = True # negrito ligado
paragrafo.style.font.italic = True # italico ligado
paragrafo.style.font.underline = True # sublinhado ligado
paragrafo.style.font.color.rgb = RGBColor(255, 0, 0) # cor escolhida

# criaremos um novo paragrafo, porem o estilo criado anteriormente somente sera aplicado a esse paragrafo caso passemos seu nome no segundo argumento da funcao
paragrafo = documento.add_paragraph("PS: A quantidade de produtos vendidos foi de 10","EstiloLira")



### usando um estilo que já existe no word

# vendo os estilos diponiveis no word:
for estilo in documento.styles:
    print(estilo)

# adicionando paragrafo com estilo diferente
paragrafo = documento.add_paragraph("testando estilos aleatorios","macro")

# implementando tudo que fizemos no Python no texto.docx

documento.save('Texto.docx')




### Aproveitando um template do word, assim nao precisa criar o estilo no python, a criacao sera no word, depois salvaremos um arquivo (no caso o template.docx) e iremos aproveita-lo abrindo o Document com ele

template = Document("template.docx")

paragrafo = template.add_paragraph("Contrato de Prestação de Serviço", "novoEstiloTeste")
template.save("NovoArquivo.docx")




### Adicionar texto com edição dentro do texto

# criando o documento no Python
documento = Document()

faturamento = 1000

# aqui você edita tudo o que você quer
texto = f"""Fala Lira,

O faturamento da empresa ontem foi de """
paragrafo = documento.add_paragraph(texto) 

# adiciona um texto no mesmo paragrafo com o estilo em negrito, poderia tambem passar o estilo especifico como segundo argumento da funcao add_run
paragrafo.add_run(f"R${faturamento}").bold = True

texto_final = "Tamo junto, abs."
paragrafo = documento.add_paragraph(texto_final)



documento.save("Texto2.docx")




### Controle de Margem e Seções

# para alinhar o paragrafo (esquerda, direita, justificado ou centralizado)
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# criando o documento no Python
documento = Document()


# as margens do documento passamos em centimetros com a funcao Cm
for secao in documento.sections:
    secao.top_margin = Cm(0.5)
    secao.bottom_margin = Cm(1)
    secao.left_margin = Cm(1)
    secao.right_margin = Cm(1)


faturamento = 1000

# aqui você edita tudo o que você quer
texto = f"""Fala Lira,

O faturamento da empresa ontem foi de R${faturamento}

Tamo junto, abs.
"""

paragrafo = documento.add_paragraph(texto) 

# alinha o paragrafo centralizado
paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

documento.save("Texto3.docx")




### Inserir Imagem

# criando o documento no Python
documento = Document()

for secao in documento.sections:
    secao.top_margin = Cm(0.5)
    secao.bottom_margin = Cm(1)
    secao.left_margin = Cm(1)
    secao.right_margin = Cm(1)


faturamento = 1000

# aqui você edita tudo o que você quer
texto = f"""Fala Lira,

O faturamento da empresa ontem foi de R${faturamento}

Tamo junto, abs.
"""

paragrafo = documento.add_paragraph(texto)

# adiciona a imagem com o valor da largura e altura
imagem = documento.add_picture("imagem.png", width=Cm(4), height=Cm(4))

documento.save("TextoComImagem.docx")




### Inserir Tabela

documento = Document()

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# observe como inserimos o estilo por meio do argumento nomeado style, podemos ver a lista de estilos disponiveis com o loop abaixo
table = documento.add_table(rows=1, cols=3, style="Light Grid Accent 1")
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
    
    
documento.save("Tabela.docx")


### Para verificar todos os estilos de um arquivo:

# vai mostrar todos os estilos disponiveis para documento que foi o ultimo arquivo criado com Document()
for estilo in documento.styles:
    print(estilo)




### Adaptar um documento

from datetime import datetime

# utilizaremos um template pois o python vai identificar os campos a serem substituidos
contrato = Document("Contrato.docx")

# faremos um exemplo simplificado no qual os valores abaixos virao de hard code, mas normalmente podem vir de variaveis, banco de dados ou planilha entre outros
nome = "Lira da Hashtag"
item1 = "Serviço de Treinamento em Excel"
item2 = "Apostila Completa de Excel"
item3 = "Serviço de Treinamentos de Python"

# agora faremos um par chave valor com um dicionario onde o que está preenchido no modelo (template) recebera os valores definidos nas variaveis
dicionario_valores = {
    "XXXX": nome,
    "YYYY": item1,
    "ZZZZ": item2,
    "WWWW": item3,
    "DD": str(datetime.now().day),
    "MM": str(datetime.now().month),
    "AAAA": str(datetime.now().year),
}

print(datetime.now())


# tem alguma função replace? -> é executada por parágrafo
# percorrer as linhas
    # se a linha tiver o texto xxxx substitui pelo nome
for paragrafo in contrato.paragraphs:
    # para cada placeholder do dicionario
    for codigo in dicionario_valores:
        if codigo in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(codigo, dicionario_valores[codigo])

contrato.save(f"Contrato Atualizado - {nome}.docx")